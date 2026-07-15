"""Build the Grand World completion-audit Word report from its split Markdown roadmap.

The output uses the Documents skill's ``standard_business_brief`` preset:
US Letter, 1-inch margins, Calibri 11 pt, fixed 9360-DXA tables, real Word
numbering, and a memo-style first-page masthead.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SOURCE_DIR = ROOT / "docs" / "roadmap" / "completion_audit"
OUTPUT = SOURCE_DIR / "GRAND_WORLD_COMPLETION_AUDIT_AND_REVISED_ROADMAP.docx"

SOURCES = [
    SOURCE_DIR / "README.md",
    SOURCE_DIR / "00_STATUS_BASELINE.md",
    SOURCE_DIR / "01_MISSING_CORE_PILLARS.md",
    SOURCE_DIR / "02_GLOBAL_HISTORICAL_CONTENT.md",
    SOURCE_DIR / "03_EXISTING_SYSTEM_DEPTH.md",
    SOURCE_DIR / "04_PRESENTATION_UI_AUDIO.md",
    SOURCE_DIR / "05_RELEASE_READINESS.md",
    SOURCE_DIR / "06_DELIVERY_SEQUENCE_AND_GATES.md",
    SOURCE_DIR / "07_METRICS_AND_EVIDENCE.md",
]

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
MUTED = "5B6673"
LIGHT_GRAY = "F2F4F7"
BLUE_GRAY = "E8EEF5"
CALLOUT = "F4F6F9"
GRID = "C9D1DB"
WHITE = "FFFFFF"
GREEN = "1F5E3B"
GOLD = "7A5A00"
RED = "9B1C1C"

TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != TABLE_WIDTH_DXA:
        raise ValueError(f"Column widths must sum to {TABLE_WIDTH_DXA}: {widths}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), GRID)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        prevent_row_split(row)
        for index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")


def set_run_font(run, name="Calibri", size=11, color=None, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def paragraph_bottom_rule(paragraph, color=BLUE, size=10, space=5) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, separate, text, end))
    set_run_font(run, size=9, color=MUTED)


def add_hyperlink(paragraph, text: str, url: str, bold=False) -> None:
    part = paragraph.part
    rel_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "22")
    r_pr.extend((fonts, color, underline, size))
    if bold:
        r_pr.append(OxmlElement("w:b"))
    run.append(r_pr)
    txt = OxmlElement("w:t")
    txt.text = text
    run.append(txt)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_RE = re.compile(r"(\*\*.+?\*\*|`.+?`|\[[^\]]+\]\([^)]+\))")


def add_inline(paragraph, text: str, *, size=11, color=None, bold=False) -> None:
    position = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position : match.start()])
            set_run_font(run, size=size, color=color, bold=bold)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, color=color, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Courier New", size=max(size - 1, 8), color=NAVY)
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), LIGHT_GRAY)
            run._element.get_or_add_rPr().append(shd)
        else:
            link = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if link:
                target = link.group(2)
                if not re.match(r"^[a-z]+://", target):
                    target = (SOURCE_DIR / target).resolve().as_uri()
                add_hyperlink(paragraph, link.group(1), target, bold=bold)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run, size=size, color=color, bold=bold)


def style_paragraph(paragraph, after=6, before=0, line=1.10) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.widow_control = True


def define_numbering(document: Document) -> tuple[int, int]:
    numbering = document.part.numbering_part.element
    existing_abstract = [int(el.get(qn("w:abstractNumId"))) for el in numbering.findall(qn("w:abstractNum"))]
    bullet_abs = max(existing_abstract, default=0) + 1
    decimal_abs = bullet_abs + 1

    def make_abstract(abs_id: int, kind: str) -> None:
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abs_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "720")
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "720")
        ind.set(qn("w:hanging"), "360")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "160")
        spacing.set(qn("w:line"), "280")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.extend((tabs, ind, spacing))
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Calibri")
        fonts.set(qn("w:hAnsi"), "Calibri")
        r_pr.append(fonts)
        lvl.extend((start, num_fmt, lvl_text, lvl_jc, p_pr, r_pr))
        abstract.append(lvl)
        numbering.append(abstract)

    make_abstract(bullet_abs, "bullet")
    make_abstract(decimal_abs, "decimal")
    return bullet_abs, decimal_abs


def new_num_id(document: Document, abstract_id: int) -> int:
    numbering = document.part.numbering_part.element
    existing = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    num_id = max(existing, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), str(abstract_id))
    num.append(abstract)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend((ilvl, num))
    p_pr.append(num_pr)


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    normal.paragraph_format.widow_control = True

    for name, size, colour, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(colour)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    table_text = styles.add_style("Audit Table Text", 1)
    table_text.font.name = "Calibri"
    table_text._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    table_text._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    table_text.font.size = Pt(9)
    table_text.paragraph_format.space_before = Pt(0)
    table_text.paragraph_format.space_after = Pt(0)
    table_text.paragraph_format.line_spacing = 1.05

    code = styles.add_style("Audit Code", 1)
    code.font.name = "Courier New"
    code._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
    code._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
    code.font.size = Pt(8.5)
    code.font.color.rgb = RGBColor.from_string(NAVY)
    code.paragraph_format.left_indent = Inches(0.15)
    code.paragraph_format.right_indent = Inches(0.15)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(6)
    code.paragraph_format.line_spacing = 1.0


def configure_page(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_paragraph(p, after=0, line=1.0)
    run = p.add_run("GRAND WORLD  |  COMPLETION AUDIT AND REVISED ROADMAP")
    set_run_font(run, size=8.5, color=MUTED, bold=True)

    first_header = section.first_page_header
    first_header.paragraphs[0].text = ""

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_paragraph(p, after=0, line=1.0)
    run = p.add_run("15 July 2026  •  Page ")
    set_run_font(run, size=8.5, color=MUTED)
    add_page_field(p)

    first_footer = section.first_page_footer
    p = first_footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(p, after=0, line=1.0)
    run = p.add_run("Grand World production-planning report  •  15 July 2026")
    set_run_font(run, size=8.5, color=MUTED)


def add_cover(document: Document) -> None:
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(42)

    kicker = document.add_paragraph()
    style_paragraph(kicker, after=8, line=1.0)
    run = kicker.add_run("PRODUCTION STATUS • 1444–1700")
    set_run_font(run, size=10, color=BLUE, bold=True)

    title = document.add_paragraph()
    style_paragraph(title, after=8, line=1.0)
    run = title.add_run("Grand World Completion Audit")
    set_run_font(run, size=28, color=NAVY, bold=True)

    subtitle = document.add_paragraph()
    style_paragraph(subtitle, after=20, line=1.10)
    run = subtitle.add_run("Revised Global Production Roadmap")
    set_run_font(run, size=16, color=DARK_BLUE)

    meta = [
        ("Assessment date", "15 July 2026"),
        ("Campaign scope", "11 November 1444 to 1 January 1700"),
        ("Current position", "Broad systems prototype with an Iberian-focused vertical-slice foundation"),
        ("Planning estimate", "25–35% of a full EU4-scale 1.0; 20–25% release-ready"),
    ]
    for label, value in meta:
        p = document.add_paragraph()
        style_paragraph(p, after=3, line=1.05)
        r = p.add_run(f"{label}: ")
        set_run_font(r, size=10.5, color=NAVY, bold=True)
        r = p.add_run(value)
        set_run_font(r, size=10.5, color=MUTED)

    rule = document.add_paragraph()
    rule.paragraph_format.space_before = Pt(12)
    rule.paragraph_format.space_after = Pt(18)
    paragraph_bottom_rule(rule, color=BLUE, size=12, space=8)

    table = document.add_table(rows=2, cols=3)
    widths = [3120, 3120, 3120]
    headings = ["Architecture", "Iberian slice", "Global product"]
    values = ["80–85%", "65–75%", "25–35%"]
    for col, text in enumerate(headings):
        cell = table.cell(0, col)
        cell.text = ""
        set_cell_shading(cell, NAVY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline(p, text, size=9, color=WHITE, bold=True)
    for col, text in enumerate(values):
        cell = table.cell(1, col)
        cell.text = ""
        set_cell_shading(cell, BLUE_GRAY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline(p, text, size=17, color=NAVY, bold=True)
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, widths)

    p = document.add_paragraph()
    style_paragraph(p, before=18, after=0, line=1.15)
    run = p.add_run(
        "The immediate production risk is scaling worldwide content before naval, colonisation, "
        "trade, HRE and Reformation schemas are locked. This report inserts those gates before "
        "unrestricted Phase 9 content production."
    )
    set_run_font(run, size=11, color=NAVY, bold=True)

    document.add_page_break()


def add_contents(document: Document) -> None:
    p = document.add_paragraph("Document Map", style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    intro = document.add_paragraph()
    add_inline(
        intro,
        "The report mirrors the repository’s split Markdown roadmap. Each workstream can be owned, reviewed and updated independently.",
    )
    items = [
        ("Executive assessment", "Completion estimates, priorities and the roadmap correction."),
        ("Status baseline", "Implemented foundations and the meaning of each completion label."),
        ("Missing core pillars", "Naval, exploration/colonisation, trade, HRE and Reformation."),
        ("Global historical content", "Country-completeness definition, research pipeline and geographic waves."),
        ("Existing-system depth", "Warfare, diplomacy, estates, government, economy and technology."),
        ("Presentation, UI and audio", "Final map, interface, portraits, rivers, sound and accessibility."),
        ("Release readiness", "AI, balance, tutorial, QA, performance, legal and compatibility."),
        ("Delivery sequence and gates", "G0–G11 production milestones and dependencies."),
        ("Metrics and evidence", "Content counts, tests, performance and release blockers."),
    ]
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Section"
    table.cell(0, 1).text = "Purpose"
    for row in items:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = ""
            p = cells[idx].paragraphs[0]
            p.style = document.styles["Audit Table Text"]
            add_inline(p, text, size=9, bold=(idx == 0))
    for cell in table.rows[0].cells:
        set_cell_shading(cell, LIGHT_GRAY)
        for p in cell.paragraphs:
            p.style = document.styles["Audit Table Text"]
            for run in p.runs:
                set_run_font(run, size=9, color=NAVY, bold=True)
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, [2700, 6660])

    p = document.add_paragraph()
    style_paragraph(p, before=16, after=6, line=1.10)
    r = p.add_run("Decision rule")
    set_run_font(r, size=11, color=NAVY, bold=True)
    p = document.add_paragraph()
    add_inline(
        p,
        "Do not call the project a global Alpha until all mandatory 1.0 pillars exist, worldwide AI can use them, and deterministic save/replay gates pass with every system enabled.",
    )
    document.add_page_break()


def table_widths(column_count: int) -> list[int]:
    if column_count == 1:
        return [9360]
    if column_count == 2:
        return [2700, 6660]
    if column_count == 3:
        return [2200, 1800, 5360]
    if column_count == 4:
        return [1750, 1450, 2750, 3410]
    base = TABLE_WIDTH_DXA // column_count
    widths = [base] * column_count
    widths[-1] += TABLE_WIDTH_DXA - sum(widths)
    return widths


def parse_table_row(line: str) -> list[str]:
    return [part.strip() for part in line.strip().strip("|").split("|")]


def is_separator_row(cells: list[str]) -> bool:
    return all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells)


def add_markdown_table(document: Document, lines: list[str]) -> None:
    rows = [parse_table_row(line) for line in lines]
    rows = [row for row in rows if not is_separator_row(row)]
    if not rows:
        return
    columns = len(rows[0])
    rows = [row[:columns] + [""] * max(0, columns - len(row)) for row in rows]
    table = document.add_table(rows=len(rows), cols=columns)
    for row_index, values in enumerate(rows):
        row = table.rows[row_index]
        for col_index, value in enumerate(values):
            cell = row.cells[col_index]
            cell.text = ""
            if row_index == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            p = cell.paragraphs[0]
            p.style = document.styles["Audit Table Text"]
            add_inline(p, value, size=9, color=NAVY if row_index == 0 else None, bold=row_index == 0)
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, table_widths(columns))
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def add_code_block(document: Document, lines: list[str]) -> None:
    p = document.add_paragraph(style="Audit Code")
    run = p.add_run("\n".join(lines))
    set_run_font(run, name="Courier New", size=8.5, color=NAVY)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), CALLOUT)
    p._p.get_or_add_pPr().append(shd)


def add_markdown(document: Document, path: Path, bullet_abs: int, decimal_abs: int, first: bool) -> None:
    start_on_new_page = not first
    lines = path.read_text(encoding="utf-8").splitlines()
    index = 0
    in_code = False
    code_lines: list[str] = []
    current_list_kind = None
    current_num_id = None

    while index < len(lines):
        raw = lines[index].rstrip()
        stripped = raw.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(document, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(raw)
            index += 1
            continue
        if not stripped or stripped == "---":
            current_list_kind = None
            current_num_id = None
            index += 1
            continue
        if stripped.startswith("|"):
            block = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                block.append(lines[index].strip())
                index += 1
            add_markdown_table(document, block)
            current_list_kind = None
            current_num_id = None
            continue

        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1))
            if path.name == "README.md" and heading.group(2) == "Companion Document":
                break
            p = document.add_paragraph(style=f"Heading {level}")
            if start_on_new_page:
                p.paragraph_format.page_break_before = True
                start_on_new_page = False
            add_inline(p, heading.group(2), size={1: 16, 2: 13, 3: 12}[level], color=BLUE if level < 3 else DARK_BLUE, bold=True)
            current_list_kind = None
            current_num_id = None
            index += 1
            continue

        bullet = re.match(r"^-\s+(.+)$", stripped)
        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if bullet or numbered:
            kind = "bullet" if bullet else "decimal"
            if kind != current_list_kind:
                current_list_kind = kind
                current_num_id = new_num_id(document, bullet_abs if kind == "bullet" else decimal_abs)
            p = document.add_paragraph()
            style_paragraph(p, after=8, line=1.167)
            apply_num(p, current_num_id)
            add_inline(p, (bullet or numbered).group(1))
            index += 1
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            nxt = lines[index].strip()
            if (
                not nxt
                or nxt.startswith("#")
                or nxt.startswith("|")
                or nxt.startswith("```")
                or re.match(r"^-\s+", nxt)
                or re.match(r"^\d+\.\s+", nxt)
                or nxt == "---"
            ):
                break
            paragraph_lines.append(nxt)
            index += 1
        p = document.add_paragraph()
        add_inline(p, " ".join(paragraph_lines))
        current_list_kind = None
        current_num_id = None


def set_document_properties(document: Document) -> None:
    props = document.core_properties
    props.title = "Grand World Completion Audit and Revised Roadmap"
    props.subject = "Production status, missing systems, content scale and release roadmap"
    props.author = "Grand World Project"
    props.keywords = "Grand World, roadmap, completion audit, grand strategy, Godot"
    props.comments = "Generated from the split Markdown roadmap in docs/roadmap/completion_audit."


def build() -> Path:
    for source in SOURCES:
        if not source.exists():
            raise FileNotFoundError(source)

    document = Document()
    configure_styles(document)
    configure_page(document)
    set_document_properties(document)
    bullet_abs, decimal_abs = define_numbering(document)

    add_cover(document)
    add_contents(document)

    for index, source in enumerate(SOURCES):
        add_markdown(document, source, bullet_abs, decimal_abs, first=index == 0)

    document.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
